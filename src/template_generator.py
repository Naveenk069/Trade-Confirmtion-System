"""
template_generator.py — Generates realistic dummy trade confirmation templates.

In a real project this module is replaced by loading actual .docx files from
your document management system. The metadata schema stays identical.
"""

import sys
from pathlib import Path
sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from src.config import TEMPLATES_DIR


# ─── Sample template definitions ─────────────────────────────────────────
TEMPLATES = [
    {
        "filename": "IRS_GoldmanSachs_USD_v1.docx",
        "metadata": {
            "trade_type": "Interest Rate Swap",
            "counterparty": "Goldman Sachs",
            "jurisdiction": "US",
            "product": "IRS",
            "version": "1.0",
            "status": "active",
        },
        "content": {
            "title": "INTEREST RATE SWAP CONFIRMATION",
            "date": "March 1, 2026",
            "parties": "Goldman Sachs Bank USA (Party A) and ABC Capital LLC (Party B)",
            "notional": "USD 50,000,000",
            "fixed_rate": "4.25% per annum",
            "floating_rate": "SOFR + 100bps",
            "payment_freq": "Semi-Annual",
            "start_date": "March 15, 2026",
            "maturity": "March 15, 2031",
            "day_count": "Actual/360",
            "governing_law": "New York",
            "isda_version": "2002 ISDA Master Agreement",
            "body": (
                "This Confirmation supplements, forms part of, and is subject to the ISDA Master Agreement "
                "dated as of the Trade Date between the parties. The fixed rate payer shall pay the Fixed Amount "
                "on each Payment Date calculated on the Notional Amount at the Fixed Rate. The floating rate payer "
                "shall pay the Floating Amount determined by reference to SOFR published by the Federal Reserve Bank "
                "of New York. All payments shall be made in immediately available funds. Early termination events "
                "include failure to pay, bankruptcy, and cross-default provisions as defined in the Master Agreement. "
                "Credit Support Annex (CSA) governs collateral posting obligations. Variation margin shall be "
                "exchanged daily on a net basis. Initial margin requirements apply per SIMM methodology."
            ),
        }
    },
    {
        "filename": "FX_Forward_JPMorgan_EUR_v1.docx",
        "metadata": {
            "trade_type": "FX Forward",
            "counterparty": "JP Morgan",
            "jurisdiction": "UK",
            "product": "FX",
            "version": "1.0",
            "status": "active",
        },
        "content": {
            "title": "FOREIGN EXCHANGE FORWARD CONFIRMATION",
            "date": "March 1, 2026",
            "parties": "J.P. Morgan Securities plc (Seller) and ABC Capital LLC (Buyer)",
            "notional": "EUR 20,000,000",
            "fixed_rate": "N/A",
            "floating_rate": "N/A",
            "payment_freq": "At Maturity",
            "start_date": "March 5, 2026",
            "maturity": "June 5, 2026",
            "day_count": "Actual/365",
            "governing_law": "English",
            "isda_version": "2002 ISDA Master Agreement",
            "body": (
                "This Confirmation confirms the terms of the FX Forward Transaction entered into between the parties. "
                "The Seller agrees to sell and the Buyer agrees to purchase EUR 20,000,000 against USD at the agreed "
                "forward exchange rate of 1.0850 USD per EUR on the Settlement Date. Settlement shall be by physical "
                "delivery of the Currency Amounts. Both parties shall make payment in same-day funds via SWIFT. "
                "The transaction is subject to the FX and Currency Option Definitions (2005 ISDA). Netting applies "
                "under the applicable ISDA Master Agreement. This confirmation is governed by English law and subject "
                "to the exclusive jurisdiction of the English courts."
            ),
        }
    },
    {
        "filename": "CDS_Barclays_EUR_v1.docx",
        "metadata": {
            "trade_type": "Credit Default Swap",
            "counterparty": "Barclays",
            "jurisdiction": "EU",
            "product": "CDS",
            "version": "1.0",
            "status": "active",
        },
        "content": {
            "title": "CREDIT DEFAULT SWAP CONFIRMATION",
            "date": "February 15, 2026",
            "parties": "Barclays Bank PLC (Protection Seller) and ABC Capital LLC (Protection Buyer)",
            "notional": "EUR 10,000,000",
            "fixed_rate": "120bps per annum (CDS Spread)",
            "floating_rate": "N/A",
            "payment_freq": "Quarterly",
            "start_date": "February 20, 2026",
            "maturity": "February 20, 2031",
            "day_count": "Actual/360",
            "governing_law": "English",
            "isda_version": "2003 ISDA Credit Derivatives Definitions",
            "body": (
                "This Confirmation is subject to the 2003 ISDA Credit Derivatives Definitions. The Reference Entity "
                "is European Investment Bank. Credit Events include Bankruptcy, Failure to Pay, and Restructuring "
                "as defined under the Definitions. Upon occurrence of a Credit Event, the Protection Seller shall "
                "pay the Credit Event Payment to the Protection Buyer. Physical Settlement applies; the Deliverable "
                "Obligations must meet the Deliverable Obligation Characteristics. The CDS spread of 120 basis points "
                "per annum is payable quarterly by the Protection Buyer. EMIR clearing obligations apply, and the "
                "trade is subject to mandatory reporting under EMIR Refit."
            ),
        }
    },
    {
        "filename": "IRS_Citibank_GBP_v2.docx",
        "metadata": {
            "trade_type": "Interest Rate Swap",
            "counterparty": "Citibank",
            "jurisdiction": "UK",
            "product": "IRS",
            "version": "2.0",
            "status": "active",
        },
        "content": {
            "title": "INTEREST RATE SWAP CONFIRMATION",
            "date": "January 10, 2026",
            "parties": "Citibank N.A. London Branch (Party A) and ABC Capital LLC (Party B)",
            "notional": "GBP 30,000,000",
            "fixed_rate": "3.95% per annum",
            "floating_rate": "SONIA + 75bps",
            "payment_freq": "Quarterly",
            "start_date": "January 20, 2026",
            "maturity": "January 20, 2029",
            "day_count": "Actual/365",
            "governing_law": "English",
            "isda_version": "2002 ISDA Master Agreement",
            "body": (
                "This Confirmation supplements and forms part of the 2002 ISDA Master Agreement. The Transaction "
                "is a Sterling Interest Rate Swap referencing SONIA as the overnight risk-free rate published by "
                "the Bank of England. Compounded SONIA in arrears applies for the floating leg calculation. "
                "The lookback period is 5 business days. Payment netting applies on the same payment dates. "
                "The transaction is centrally cleared via LCH SwapClear. Initial and variation margin requirements "
                "apply per LCH rulebook. MiFID II transaction reporting obligations apply to both parties."
            ),
        }
    },
    {
        "filename": "EquitySwap_DeutscheBank_US_v1.docx",
        "metadata": {
            "trade_type": "Equity Swap",
            "counterparty": "Deutsche Bank",
            "jurisdiction": "US",
            "product": "Equity Swap",
            "version": "1.0",
            "status": "active",
        },
        "content": {
            "title": "EQUITY SWAP CONFIRMATION",
            "date": "February 28, 2026",
            "parties": "Deutsche Bank AG New York Branch (Equity Amount Payer) and ABC Capital LLC (Fixed Amount Payer)",
            "notional": "USD 15,000,000",
            "fixed_rate": "SOFR + 50bps",
            "floating_rate": "Total Return on S&P 500 Index",
            "payment_freq": "Quarterly",
            "start_date": "March 1, 2026",
            "maturity": "March 1, 2028",
            "day_count": "Actual/360",
            "governing_law": "New York",
            "isda_version": "2002 ISDA Master Agreement + 2011 Equity Derivatives Definitions",
            "body": (
                "This Confirmation is subject to the 2002 ISDA Master Agreement and the 2011 ISDA Equity Derivatives "
                "Definitions. The Equity Amount Payer shall pay amounts based on the total return of the S&P 500 Index, "
                "including price appreciation and dividend equivalent amounts. The Fixed Amount Payer shall pay SOFR + 50bps "
                "on the Notional Amount. Dividend adjustment applies on ex-dividend dates. Index disruption events and "
                "adjustment provisions apply per the Definitions. The transaction is subject to Dodd-Frank reporting "
                "and clearing requirements. Swap Data Repository reporting to the CFTC is required within T+1."
            ),
        }
    },
    {
        "filename": "FX_Forward_HSBC_JPY_v1.docx",
        "metadata": {
            "trade_type": "FX Forward",
            "counterparty": "HSBC",
            "jurisdiction": "UK",
            "product": "FX",
            "version": "1.0",
            "status": "active",
        },
        "content": {
            "title": "FOREIGN EXCHANGE FORWARD CONFIRMATION",
            "date": "March 3, 2026",
            "parties": "HSBC Bank plc (Seller) and ABC Capital LLC (Buyer)",
            "notional": "JPY 2,000,000,000",
            "fixed_rate": "N/A",
            "floating_rate": "N/A",
            "payment_freq": "At Maturity",
            "start_date": "March 10, 2026",
            "maturity": "September 10, 2026",
            "day_count": "Actual/365",
            "governing_law": "English",
            "isda_version": "2002 ISDA Master Agreement",
            "body": (
                "This Confirmation confirms the FX Forward Transaction for the purchase and sale of Japanese Yen. "
                "The Seller agrees to deliver JPY 2,000,000,000 and the Buyer agrees to deliver USD at an agreed "
                "forward rate of 149.50 JPY per USD on the Settlement Date. SWIFT payment instructions apply. "
                "The Bank of Japan Tokyo fixing rate serves as the reference rate for settlement disputes. "
                "This transaction is subject to FCA regulation and EMIR reporting requirements."
            ),
        }
    },
    {
        "filename": "IRS_GoldmanSachs_USD_v2.docx",
        "metadata": {
            "trade_type": "Interest Rate Swap",
            "counterparty": "Goldman Sachs",
            "jurisdiction": "US",
            "product": "IRS",
            "version": "2.0",
            "status": "active",
        },
        "content": {
            "title": "INTEREST RATE SWAP CONFIRMATION (AMENDED)",
            "date": "March 5, 2026",
            "parties": "Goldman Sachs Bank USA (Party A) and ABC Capital LLC (Party B)",
            "notional": "USD 75,000,000",
            "fixed_rate": "4.50% per annum",
            "floating_rate": "SOFR + 120bps",
            "payment_freq": "Quarterly",
            "start_date": "March 20, 2026",
            "maturity": "March 20, 2033",
            "day_count": "Actual/360",
            "governing_law": "New York",
            "isda_version": "2002 ISDA Master Agreement",
            "body": (
                "This amended Confirmation supersedes and replaces the prior Confirmation dated March 1, 2026. "
                "The Notional Amount has been increased to USD 75,000,000. The fixed rate has been revised upward "
                "to 4.50% per annum to reflect current market conditions. Compounded SOFR in arrears with a "
                "5-business-day lookback applies. CME Clearing applies; initial margin is calculated per SIMM. "
                "All other terms remain unchanged from the original Confirmation."
            ),
        }
    },
    {
        "filename": "CDS_BNPParibas_US_v1.docx",
        "metadata": {
            "trade_type": "Credit Default Swap",
            "counterparty": "BNP Paribas",
            "jurisdiction": "US",
            "product": "CDS",
            "version": "1.0",
            "status": "draft",
        },
        "content": {
            "title": "CREDIT DEFAULT SWAP CONFIRMATION — DRAFT",
            "date": "March 7, 2026",
            "parties": "BNP Paribas Securities Corp (Protection Seller) and ABC Capital LLC (Protection Buyer)",
            "notional": "USD 5,000,000",
            "fixed_rate": "85bps per annum",
            "floating_rate": "N/A",
            "payment_freq": "Quarterly",
            "start_date": "March 15, 2026",
            "maturity": "March 15, 2029",
            "day_count": "Actual/360",
            "governing_law": "New York",
            "isda_version": "2003 ISDA Credit Derivatives Definitions",
            "body": (
                "DRAFT — Subject to legal review. This Confirmation references Ford Motor Company as the Reference Entity. "
                "Credit Events include Bankruptcy and Failure to Pay. Cash Settlement applies with a Valuation Date "
                "of 5 business days after the Credit Event Notice. The Final Price will be determined by a Dealer Poll. "
                "CFTC reporting requirements apply. SEF execution may be required pending regulatory determination."
            ),
        }
    },
]


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT


def _add_field_row(doc: Document, label: str, value: str) -> None:
    """Add a bold label + value row."""
    p = doc.add_paragraph()
    run_label = p.add_run(f"{label}: ")
    run_label.bold = True
    run_label.font.size = Pt(10)
    run_value = p.add_run(value)
    run_value.font.size = Pt(10)


def generate_template(template_def: dict, output_dir: Path) -> Path:
    """Generate a single .docx template from its definition dict."""
    doc = Document()

    # Header / title
    content = template_def["content"]
    meta = template_def["metadata"]

    _add_heading(doc, content["title"], level=1)
    doc.add_paragraph()

    # Key trade terms table-style
    _add_field_row(doc, "Date",               content["date"])
    _add_field_row(doc, "Parties",            content["parties"])
    _add_field_row(doc, "Trade Type",         meta["trade_type"])
    _add_field_row(doc, "Product",            meta["product"])
    _add_field_row(doc, "Counterparty",       meta["counterparty"])
    _add_field_row(doc, "Jurisdiction",       meta["jurisdiction"])
    _add_field_row(doc, "Governing Law",      content["governing_law"])
    _add_field_row(doc, "ISDA Agreement",     content["isda_version"])
    doc.add_paragraph()

    _add_heading(doc, "Economic Terms", level=2)
    _add_field_row(doc, "Notional Amount",    content["notional"])
    _add_field_row(doc, "Fixed Rate",         content["fixed_rate"])
    _add_field_row(doc, "Floating Rate",      content["floating_rate"])
    _add_field_row(doc, "Payment Frequency",  content["payment_freq"])
    _add_field_row(doc, "Effective Date",     content["start_date"])
    _add_field_row(doc, "Maturity Date",      content["maturity"])
    _add_field_row(doc, "Day Count",          content["day_count"])
    doc.add_paragraph()

    _add_heading(doc, "Confirmation Terms", level=2)
    doc.add_paragraph(content["body"])
    doc.add_paragraph()

    # Footer metadata (hidden from UI but embedded in doc for parsing)
    _add_heading(doc, "Document Metadata", level=2)
    _add_field_row(doc, "Version",  meta["version"])
    _add_field_row(doc, "Status",   meta["status"])

    output_path = output_dir / template_def["filename"]
    doc.save(str(output_path))
    return output_path


def generate_all_templates(output_dir: Path = TEMPLATES_DIR) -> list[Path]:
    """Generate all dummy templates and return their paths."""
    output_dir.mkdir(parents=True, exist_ok=True)
    generated = []
    for t in TEMPLATES:
        path = generate_template(t, output_dir)
        generated.append(path)
        print(f"  ✓ Generated: {path.name}")
    return generated


if __name__ == "__main__":
    print("\n📄 Generating sample trade confirmation templates...\n")
    paths = generate_all_templates()
    print(f"\n✅ {len(paths)} templates created in: {TEMPLATES_DIR}\n")
