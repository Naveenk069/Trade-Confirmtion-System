"""
template_populator.py — Pre-populated template generation (Iteration 3).

Three-stage pipeline:
  1. AUTO-FILL     — inject trade data (from JSON/CSV/dict) into template fields
  2. AI CLAUSES    — generate contextually appropriate clause text using rules + AI patterns
  3. VALIDATION    — validate all fields against approved clause library and business rules

Usage:
    populator = TemplatePopulator()

    trade_data = {
        "trade_type":   "Interest Rate Swap",
        "counterparty": "Goldman Sachs",
        "notional":     "USD 50,000,000",
        "fixed_rate":   "4.25%",
        ...
    }

    result = populator.generate(trade_data)
    result.save("output/IRS_Goldman_populated.docx")
"""

import re
import sys
import json
import csv
import copy
from dataclasses import dataclass, field
from datetime import datetime, date
from pathlib import Path
from typing import Optional, Any

sys.path.insert(0, str(Path(__file__).resolve().parent.parent.parent))

from src.config import TEMPLATES_DIR


# ─── Trade Data Model ─────────────────────────────────────────────────────

@dataclass
class TradeData:
    """
    Structured trade data input.
    All fields are optional strings — the populator handles missing values.
    """
    # Identity
    trade_type:        str = ""
    product:           str = ""
    counterparty:      str = ""
    jurisdiction:      str = ""

    # Economic terms
    notional_amount:   str = ""
    currency:          str = ""
    fixed_rate:        str = ""
    floating_rate:     str = ""
    payment_frequency: str = ""
    effective_date:    str = ""
    maturity_date:     str = ""
    day_count:         str = ""

    # Legal
    governing_law:     str = ""
    isda_version:      str = ""
    clearing_venue:    str = ""

    # Parties
    party_a:           str = ""
    party_b:           str = ""
    party_a_role:      str = ""
    party_b_role:      str = ""

    # Meta
    trade_date:        str = ""
    trader_name:       str = ""

    # Extra fields (catch-all for custom data)
    extras:            dict = field(default_factory=dict)

    @classmethod
    def from_dict(cls, data: dict) -> "TradeData":
        """Build TradeData from a flat dictionary."""
        known = {f for f in cls.__dataclass_fields__ if f != "extras"}
        init_kwargs = {k: str(v) for k, v in data.items() if k in known}
        extras = {k: v for k, v in data.items() if k not in known}
        return cls(**init_kwargs, extras=extras)

    @classmethod
    def from_json(cls, json_str: str) -> "TradeData":
        return cls.from_dict(json.loads(json_str))

    @classmethod
    def from_json_file(cls, path: Path) -> "TradeData":
        return cls.from_dict(json.loads(Path(path).read_text()))

    @classmethod
    def from_csv_row(cls, csv_path: Path, row_index: int = 0) -> "TradeData":
        """Load a single row from a CSV file."""
        with open(csv_path, newline="", encoding="utf-8") as f:
            reader = csv.DictReader(f)
            rows = list(reader)
        if not rows:
            raise ValueError(f"No rows in CSV: {csv_path}")
        return cls.from_dict(dict(rows[row_index]))

    def to_dict(self) -> dict:
        d = {k: v for k, v in self.__dict__.items() if k != "extras"}
        d.update(self.extras)
        return d


# ─── Validation Models ────────────────────────────────────────────────────

@dataclass
class ValidationIssue:
    field:    str
    severity: str   # "ERROR" | "WARNING" | "INFO"
    message:  str
    rule_id:  str


@dataclass
class PopulatedTemplate:
    """Result of the three-stage generation pipeline."""
    filename:         str
    trade_data:       TradeData
    populated_text:   str          # Full document text with fields filled
    issues:           list[ValidationIssue] = field(default_factory=list)
    autofill_count:   int = 0
    ai_clause_count:  int = 0
    validation_score: float = 0.0  # 0–100

    @property
    def errors(self) -> list[ValidationIssue]:
        return [i for i in self.issues if i.severity == "ERROR"]

    @property
    def warnings(self) -> list[ValidationIssue]:
        return [i for i in self.issues if i.severity == "WARNING"]

    @property
    def is_valid(self) -> bool:
        return len(self.errors) == 0

    def save(self, output_path: Path) -> Path:
        """Save the populated template as a .docx file."""
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        _save_as_docx(self.populated_text, self.trade_data, output_path)
        return output_path

    def summary(self) -> str:
        return (
            f"PopulatedTemplate: {self.filename}\n"
            f"  Auto-filled fields : {self.autofill_count}\n"
            f"  AI-generated clauses: {self.ai_clause_count}\n"
            f"  Validation score   : {self.validation_score:.0f}/100\n"
            f"  Errors             : {len(self.errors)}\n"
            f"  Warnings           : {len(self.warnings)}\n"
            f"  Ready for review   : {'YES' if self.is_valid else 'NO — fix errors first'}"
        )


# ─── Approved Clause Library ──────────────────────────────────────────────

APPROVED_CLAUSES = {
    # Governing law by jurisdiction
    "governing_law": {
        "US":   "New York",
        "UK":   "English",
        "EU":   "English",
        "APAC": "English",
    },
    # ISDA version by trade type
    "isda_version": {
        "Interest Rate Swap":   "2002 ISDA Master Agreement",
        "FX Forward":           "2002 ISDA Master Agreement",
        "Credit Default Swap":  "2003 ISDA Credit Derivatives Definitions",
        "Equity Swap":          "2002 ISDA Master Agreement + 2011 Equity Derivatives Definitions",
        "Total Return Swap":    "2002 ISDA Master Agreement",
    },
    # Day count by currency
    "day_count": {
        "USD": "Actual/360",
        "EUR": "Actual/360",
        "GBP": "Actual/365",
        "JPY": "Actual/360",
        "CHF": "Actual/360",
    },
    # Floating rate benchmarks by currency
    "floating_rate_benchmark": {
        "USD": "SOFR",
        "GBP": "SONIA",
        "EUR": "EURIBOR",
        "JPY": "TONA",
        "CHF": "SARON",
    },
    # Clearing venues
    "clearing_venue": {
        "Interest Rate Swap": "LCH SwapClear",
        "Credit Default Swap": "ICE Clear Credit",
        "Equity Swap": "Not cleared (bilateral)",
        "FX Forward": "Not cleared (bilateral)",
    },
}


# ─── Business Rules ───────────────────────────────────────────────────────

def _validate_date_order(effective: str, maturity: str) -> bool:
    """Return True if effective_date < maturity_date."""
    try:
        fmt = "%B %d, %Y"
        eff = datetime.strptime(effective.strip(), fmt)
        mat = datetime.strptime(maturity.strip(), fmt)
        return eff < mat
    except Exception:
        return True  # Can't parse → don't flag

def _validate_rate(rate_str: str) -> bool:
    """Return True if rate looks like a valid percentage."""
    match = re.search(r"\d+(\.\d+)?", rate_str)
    if not match:
        return False
    val = float(match.group())
    return 0 < val < 50  # Sanity: 0% < rate < 50%

def _validate_notional(notional_str: str) -> bool:
    """Return True if notional is a positive number."""
    match = re.search(r"[\d,]+", notional_str.replace(",", ""))
    if not match:
        return False
    return float(re.sub(r"[^\d.]", "", notional_str.replace(",", ""))) > 0


BUSINESS_RULES = [
    {
        "rule_id":  "BR-001",
        "severity": "ERROR",
        "field":    "governing_law",
        "check":    lambda td: bool(td.governing_law),
        "message":  "Governing law must be specified.",
    },
    {
        "rule_id":  "BR-002",
        "severity": "ERROR",
        "field":    "isda_version",
        "check":    lambda td: bool(td.isda_version),
        "message":  "ISDA agreement version must be specified.",
    },
    {
        "rule_id":  "BR-003",
        "severity": "ERROR",
        "field":    "notional_amount",
        "check":    lambda td: not td.notional_amount or _validate_notional(td.notional_amount),
        "message":  "Notional amount must be a positive number.",
    },
    {
        "rule_id":  "BR-004",
        "severity": "ERROR",
        "field":    "effective_date",
        "check":    lambda td: not (td.effective_date and td.maturity_date) or
                               _validate_date_order(td.effective_date, td.maturity_date),
        "message":  "Effective date must be before maturity date.",
    },
    {
        "rule_id":  "BR-005",
        "severity": "WARNING",
        "field":    "fixed_rate",
        "check":    lambda td: not td.fixed_rate or _validate_rate(td.fixed_rate),
        "message":  "Fixed rate appears outside normal range (0%–50%). Please verify.",
    },
    {
        "rule_id":  "BR-006",
        "severity": "WARNING",
        "field":    "counterparty",
        "check":    lambda td: bool(td.counterparty or td.party_a),
        "message":  "Counterparty or Party A name should be specified.",
    },
    {
        "rule_id":  "BR-007",
        "severity": "WARNING",
        "field":    "governing_law",
        "check":    lambda td: not (td.jurisdiction and td.governing_law) or
                               APPROVED_CLAUSES["governing_law"].get(td.jurisdiction, "").lower()
                               in td.governing_law.lower(),
        "message":  "Governing law does not match expected law for this jurisdiction.",
    },
    {
        "rule_id":  "BR-008",
        "severity": "INFO",
        "field":    "payment_frequency",
        "check":    lambda td: bool(td.payment_frequency),
        "message":  "Payment frequency not specified — defaulting to Quarterly.",
    },
]


# ─── Template Populator ───────────────────────────────────────────────────

class TemplatePopulator:
    """
    Three-stage pipeline to generate a pre-populated trade confirmation.

    Stage 1 — AUTO-FILL:
        Inject known trade data values into template placeholders.

    Stage 2 — AI CLAUSE SUGGESTIONS:
        Use rule-based clause library to fill in any remaining empty fields
        (governing law, ISDA version, day count, floating rate benchmark).

    Stage 3 — VALIDATION:
        Run all business rules and return ValidationIssues.
    """

    def generate(
        self,
        trade_data: TradeData | dict,
        template_filename: Optional[str] = None,
    ) -> PopulatedTemplate:
        """
        Generate a pre-populated trade confirmation.

        Args:
            trade_data:         Structured trade data (TradeData or dict).
            template_filename:  Specific template .docx to use as base.
                                If None, auto-selected from trade_type.

        Returns:
            PopulatedTemplate ready to save or display.
        """
        if isinstance(trade_data, dict):
            trade_data = TradeData.from_dict(trade_data)

        # Select base template
        base_text = self._select_template(trade_data, template_filename)
        output_filename = self._make_filename(trade_data)

        # Stage 1: Auto-fill from trade data
        filled_text, autofill_count = self._autofill(base_text, trade_data)

        # Stage 2: AI clause suggestions for empty fields
        enriched_data, ai_count = self._suggest_clauses(trade_data)
        filled_text, extra_fills = self._autofill(filled_text, enriched_data)
        autofill_count += extra_fills

        # Stage 3: Validation
        issues = self._validate(enriched_data)
        score  = self._compute_score(enriched_data, issues)

        return PopulatedTemplate(
            filename=output_filename,
            trade_data=enriched_data,
            populated_text=filled_text,
            issues=issues,
            autofill_count=autofill_count,
            ai_clause_count=ai_count,
            validation_score=score,
        )

    def generate_from_json(self, json_str: str) -> PopulatedTemplate:
        return self.generate(TradeData.from_json(json_str))

    def generate_from_file(self, json_path: Path) -> PopulatedTemplate:
        return self.generate(TradeData.from_json_file(json_path))

    def generate_from_csv(self, csv_path: Path, row: int = 0) -> PopulatedTemplate:
        return self.generate(TradeData.from_csv_row(csv_path, row))

    # ── Stage 1: Auto-fill ────────────────────────────────────────────

    # Template placeholder → TradeData field mapping
    _PLACEHOLDER_MAP = {
        "[TRADE_TYPE]":        "trade_type",
        "[PRODUCT]":           "product",
        "[COUNTERPARTY]":      "counterparty",
        "[JURISDICTION]":      "jurisdiction",
        "[NOTIONAL]":          "notional_amount",
        "[CURRENCY]":          "currency",
        "[FIXED_RATE]":        "fixed_rate",
        "[FLOATING_RATE]":     "floating_rate",
        "[PAYMENT_FREQUENCY]": "payment_frequency",
        "[EFFECTIVE_DATE]":    "effective_date",
        "[MATURITY_DATE]":     "maturity_date",
        "[DAY_COUNT]":         "day_count",
        "[GOVERNING_LAW]":     "governing_law",
        "[ISDA_VERSION]":      "isda_version",
        "[CLEARING_VENUE]":    "clearing_venue",
        "[PARTY_A]":           "party_a",
        "[PARTY_B]":           "party_b",
        "[PARTY_A_ROLE]":      "party_a_role",
        "[PARTY_B_ROLE]":      "party_b_role",
        "[TRADE_DATE]":        "trade_date",
        "[TRADER_NAME]":       "trader_name",
    }

    def _autofill(self, text: str, td: TradeData) -> tuple[str, int]:
        """Replace all known placeholders with trade data values."""
        count = 0
        data  = td.to_dict()

        for placeholder, field_key in self._PLACEHOLDER_MAP.items():
            value = data.get(field_key, "")
            if value and placeholder in text:
                text  = text.replace(placeholder, value)
                count += 1

        # Also replace regex-style field labels in the document
        field_map = {
            r"(Notional Amount\s*:\s*)([^\n]*)":    ("notional_amount",   td.notional_amount),
            r"(Fixed Rate\s*:\s*)([^\n]*)":          ("fixed_rate",        td.fixed_rate),
            r"(Floating Rate\s*:\s*)([^\n]*)":       ("floating_rate",     td.floating_rate),
            r"(Effective Date\s*:\s*)([^\n]*)":      ("effective_date",    td.effective_date),
            r"(Maturity Date\s*:\s*)([^\n]*)":       ("maturity_date",     td.maturity_date),
            r"(Payment Frequency\s*:\s*)([^\n]*)":   ("payment_frequency", td.payment_frequency),
            r"(Day Count\s*:\s*)([^\n]*)":           ("day_count",         td.day_count),
            r"(Governing Law\s*:\s*)([^\n]*)":       ("governing_law",     td.governing_law),
            r"(ISDA Agreement\s*:\s*)([^\n]*)":      ("isda_version",      td.isda_version),
            r"(Counterparty\s*:\s*)([^\n]*)":        ("counterparty",      td.counterparty),
            r"(Jurisdiction\s*:\s*)([^\n]*)":        ("jurisdiction",      td.jurisdiction),
            r"(Trade Type\s*:\s*)([^\n]*)":          ("trade_type",        td.trade_type),
        }

        for pattern, (key, new_val) in field_map.items():
            if new_val:
                def replacer(m, v=new_val):
                    count_holder = [0]
                    count_holder[0] += 1
                    return m.group(1) + v
                new_text = re.sub(pattern, replacer, text, flags=re.IGNORECASE)
                if new_text != text:
                    count += 1
                    text = new_text

        return text, count

    # ── Stage 2: AI Clause Suggestions ───────────────────────────────

    def _suggest_clauses(self, td: TradeData) -> tuple[TradeData, int]:
        """
        Fill in empty fields using the approved clause library.
        This is the 'AI suggestion' layer — deterministic rules + library lookup.

        In production, this can be augmented with an LLM call for
        more sophisticated clause generation.
        """
        td = copy.deepcopy(td)
        count = 0

        # Governing law from jurisdiction
        if not td.governing_law and td.jurisdiction:
            suggested = APPROVED_CLAUSES["governing_law"].get(td.jurisdiction)
            if suggested:
                td.governing_law = suggested
                count += 1

        # ISDA version from trade type
        if not td.isda_version and td.trade_type:
            suggested = APPROVED_CLAUSES["isda_version"].get(td.trade_type)
            if suggested:
                td.isda_version = suggested
                count += 1

        # Day count from currency
        if not td.day_count and td.currency:
            suggested = APPROVED_CLAUSES["day_count"].get(td.currency.upper())
            if suggested:
                td.day_count = suggested
                count += 1

        # Floating rate benchmark from currency
        if not td.floating_rate and td.currency:
            benchmark = APPROVED_CLAUSES["floating_rate_benchmark"].get(td.currency.upper())
            if benchmark:
                td.floating_rate = f"{benchmark} + [SPREAD]bps"
                count += 1

        # Clearing venue from trade type
        if not td.clearing_venue and td.trade_type:
            suggested = APPROVED_CLAUSES["clearing_venue"].get(td.trade_type)
            if suggested:
                td.clearing_venue = suggested
                count += 1

        # Payment frequency default
        if not td.payment_frequency:
            defaults = {
                "Interest Rate Swap":  "Quarterly",
                "Credit Default Swap": "Quarterly",
                "Equity Swap":         "Quarterly",
                "FX Forward":          "At Maturity",
            }
            suggested = defaults.get(td.trade_type)
            if suggested:
                td.payment_frequency = suggested
                count += 1

        # Trade date default to today
        if not td.trade_date:
            td.trade_date = datetime.now().strftime("%B %d, %Y")
            count += 1

        # Party roles from trade type
        if td.trade_type == "Credit Default Swap":
            if not td.party_a_role: td.party_a_role = "Protection Seller"; count += 1
            if not td.party_b_role: td.party_b_role = "Protection Buyer";  count += 1
        elif td.trade_type == "Equity Swap":
            if not td.party_a_role: td.party_a_role = "Equity Amount Payer"; count += 1
            if not td.party_b_role: td.party_b_role = "Fixed Amount Payer";  count += 1
        elif td.trade_type in ("Interest Rate Swap", "FX Forward"):
            if not td.party_a_role: td.party_a_role = "Party A"; count += 1
            if not td.party_b_role: td.party_b_role = "Party B"; count += 1

        return td, count

    # ── Stage 3: Validation ───────────────────────────────────────────

    def _validate(self, td: TradeData) -> list[ValidationIssue]:
        issues = []
        for rule in BUSINESS_RULES:
            try:
                passed = rule["check"](td)
            except Exception:
                passed = True  # Don't crash on validation errors

            if not passed:
                issues.append(ValidationIssue(
                    field=rule["field"],
                    severity=rule["severity"],
                    message=rule["message"],
                    rule_id=rule["rule_id"],
                ))

        return sorted(issues, key=lambda i: {"ERROR": 0, "WARNING": 1, "INFO": 2}[i.severity])

    def _compute_score(self, td: TradeData, issues: list[ValidationIssue]) -> float:
        """Compute a validation score 0–100."""
        # Required fields
        required = ["trade_type", "counterparty", "notional_amount",
                    "effective_date", "maturity_date", "governing_law", "isda_version"]
        present = sum(1 for f in required if getattr(td, f, ""))
        completeness = present / len(required) * 70  # 70 points for completeness

        # Deduct for issues
        deductions = sum({"ERROR": 20, "WARNING": 5, "INFO": 1}.get(i.severity, 0)
                         for i in issues)

        return max(0.0, min(100.0, completeness + 30 - deductions))

    # ── Template selection ────────────────────────────────────────────

    def _select_template(
        self, td: TradeData, template_filename: Optional[str]
    ) -> str:
        """Select the best base template text for this trade."""
        if template_filename:
            path = TEMPLATES_DIR / template_filename
            if path.exists():
                return self._read_template(path)

        # Auto-select by trade type
        candidates = list(TEMPLATES_DIR.glob("*.docx"))
        for candidate in sorted(candidates):
            name_lower = candidate.name.lower()
            tt_lower   = td.trade_type.lower().replace(" ", "_")
            if td.trade_type and (
                "irs" in name_lower and "swap" in td.trade_type.lower() or
                "fx" in name_lower and "fx" in td.trade_type.lower() or
                "cds" in name_lower and "credit" in td.trade_type.lower() or
                "equity" in name_lower and "equity" in td.trade_type.lower()
            ):
                return self._read_template(candidate)

        # Fallback: generic template string
        return self._generic_template()

    def _read_template(self, path: Path) -> str:
        """Read text from a .docx template."""
        try:
            from src.document_parser import DocumentParser
            parser = DocumentParser()
            doc = parser.parse(path)
            return doc.full_text
        except Exception:
            return self._generic_template()

    def _generic_template(self) -> str:
        return """TRADE CONFIRMATION

Date: [TRADE_DATE]
Parties: [PARTY_A] ([PARTY_A_ROLE]) and [PARTY_B] ([PARTY_B_ROLE])
Trade Type: [TRADE_TYPE]
Product: [PRODUCT]
Counterparty: [COUNTERPARTY]
Jurisdiction: [JURISDICTION]
Governing Law: [GOVERNING_LAW]
ISDA Agreement: [ISDA_VERSION]

Economic Terms
Notional Amount: [NOTIONAL]
Fixed Rate: [FIXED_RATE]
Floating Rate: [FLOATING_RATE]
Payment Frequency: [PAYMENT_FREQUENCY]
Effective Date: [EFFECTIVE_DATE]
Maturity Date: [MATURITY_DATE]
Day Count: [DAY_COUNT]

Clearing Venue: [CLEARING_VENUE]

This Confirmation supplements and forms part of the ISDA Master Agreement
between the parties. All terms are subject to the governing law and
jurisdiction specified above.

Version: 1.0
Status: draft
"""

    def _make_filename(self, td: TradeData) -> str:
        parts = []
        if td.trade_type:
            tt = td.trade_type.replace(" ", "")[:12]
            parts.append(tt)
        if td.counterparty:
            cp = re.sub(r"[^A-Za-z]", "", td.counterparty.split()[0])[:10]
            parts.append(cp)
        if td.currency:
            parts.append(td.currency.upper()[:3])
        parts.append("populated_v1.docx")
        return "_".join(parts) if len(parts) > 1 else "TradeConfirmation_populated_v1.docx"


# ─── .docx saver ──────────────────────────────────────────────────────────

def _save_as_docx(text: str, td: TradeData, output_path: Path) -> None:
    """Save populated text as a formatted .docx file."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    lines = text.split("\n")
    for line in lines:
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue

        # Title (ALL CAPS first line)
        if line == lines[0].strip() and line.isupper():
            p = doc.add_heading(line, level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Section headers
        elif line in ("Economic Terms", "Confirmation Terms", "Document Metadata"):
            doc.add_heading(line, level=2)
        # Field lines (Label: Value)
        elif ":" in line:
            p = doc.add_paragraph()
            key, _, val = line.partition(":")
            run_k = p.add_run(key.strip() + ": ")
            run_k.bold = True
            run_k.font.size = Pt(10)
            run_v = p.add_run(val.strip())
            run_v.font.size = Pt(10)
        else:
            p = doc.add_paragraph(line)
            p.runs[0].font.size = Pt(10) if p.runs else None

    doc.save(str(output_path))


if __name__ == "__main__":
    from src.template_generator import generate_all_templates
    generate_all_templates()

    populator = TemplatePopulator()

    # ── Example 1: IRS from dict ───────────────────────────────────
    print("\n" + "="*60)
    print("  ITERATION 3 — PRE-POPULATED TEMPLATE DEMO")
    print("="*60)

    irs_trade = {
        "trade_type":      "Interest Rate Swap",
        "counterparty":    "Goldman Sachs",
        "currency":        "USD",
        "notional_amount": "USD 75,000,000",
        "fixed_rate":      "4.75% per annum",
        "effective_date":  "April 1, 2026",
        "maturity_date":   "April 1, 2031",
        "jurisdiction":    "US",
        "party_a":         "Goldman Sachs Bank USA",
        "party_b":         "ABC Capital LLC",
    }

    result = populator.generate(irs_trade)
    print(f"\n{result.summary()}")

    if result.issues:
        print("\n  Validation Issues:")
        for issue in result.issues:
            icon = {"ERROR": "❌", "WARNING": "⚠", "INFO": "ℹ"}[issue.severity]
            print(f"  {icon} [{issue.rule_id}] {issue.field}: {issue.message}")

    # ── Example 2: FX Forward from JSON ───────────────────────────
    fx_json = json.dumps({
        "trade_type":      "FX Forward",
        "counterparty":    "HSBC",
        "currency":        "GBP",
        "notional_amount": "GBP 10,000,000",
        "effective_date":  "April 5, 2026",
        "maturity_date":   "October 5, 2026",
        "jurisdiction":    "UK",
        "party_a":         "HSBC Bank plc",
        "party_b":         "ABC Capital LLC",
    })

    result2 = populator.generate_from_json(fx_json)
    print(f"\n{result2.summary()}")

    # ── Save output ────────────────────────────────────────────────
    out_dir = Path("data/populated")
    out_dir.mkdir(parents=True, exist_ok=True)
    saved = result.save(out_dir / result.filename)
    print(f"\n  ✅ Saved: {saved}")
