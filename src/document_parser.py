"""
document_parser.py — Extracts structured text and metadata from .docx (and .txt) files.

Design principles:
  - Returns a clean ParsedDocument dataclass for every file
  - All downstream code works with ParsedDocument; it never touches raw files
  - Swappable: add PDF support by adding a new _parse_pdf() method
"""

import re
import sys
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from docx import Document as DocxDocument
from src.config import SUPPORTED_EXTENSIONS


# ─── Data Model ──────────────────────────────────────────────────────────

@dataclass
class ParsedDocument:
    """Represents a fully parsed trade confirmation template."""

    # Identity
    file_path: str
    filename: str

    # Full text (used for embedding)
    full_text: str

    # Extracted metadata (used for filtering)
    trade_type: str   = "unknown"
    counterparty: str = "unknown"
    jurisdiction: str = "unknown"
    product: str      = "unknown"
    version: str      = "1.0"
    status: str       = "active"

    # Structured sections (for display)
    title: str              = ""
    economic_terms: dict    = field(default_factory=dict)
    confirmation_body: str  = ""

    # Diagnostics
    parse_warnings: list    = field(default_factory=list)

    @property
    def metadata_dict(self) -> dict:
        """Returns metadata as a flat dict (used by vector store)."""
        return {
            "trade_type":   self.trade_type,
            "counterparty": self.counterparty,
            "jurisdiction": self.jurisdiction,
            "product":      self.product,
            "version":      self.version,
            "status":       self.status,
            "filename":     self.filename,
            "title":        self.title,
        }

    def __repr__(self):
        return (
            f"ParsedDocument(filename={self.filename!r}, "
            f"trade_type={self.trade_type!r}, "
            f"counterparty={self.counterparty!r}, "
            f"jurisdiction={self.jurisdiction!r})"
        )


# ─── Parser ──────────────────────────────────────────────────────────────

class DocumentParser:
    """
    Parses trade confirmation documents into ParsedDocument objects.

    Usage:
        parser = DocumentParser()
        doc = parser.parse(Path("templates/IRS_Goldman_USD.docx"))
    """

    # Field label patterns to extract from document body text
    _FIELD_PATTERNS = {
        "trade_type":   r"Trade Type\s*:\s*(.+)",
        "counterparty": r"(?:Counterparty|Parties)\s*:\s*(.+)",
        "jurisdiction": r"Jurisdiction\s*:\s*(.+)",
        "product":      r"Product\s*:\s*(.+)",
        "version":      r"Version\s*:\s*(.+)",
        "status":       r"Status\s*:\s*(.+)",
        "governing_law":r"Governing Law\s*:\s*(.+)",
        "notional":     r"Notional Amount\s*:\s*(.+)",
        "fixed_rate":   r"Fixed Rate\s*:\s*(.+)",
        "floating_rate":r"Floating Rate\s*:\s*(.+)",
        "effective_date":r"Effective Date\s*:\s*(.+)",
        "maturity_date":r"Maturity Date\s*:\s*(.+)",
        "payment_freq": r"Payment Frequency\s*:\s*(.+)",
        "day_count":    r"Day Count\s*:\s*(.+)",
        "isda_version": r"ISDA Agreement\s*:\s*(.+)",
    }

    def parse(self, file_path: Path) -> ParsedDocument:
        """
        Parse a document file and return a ParsedDocument.
        Dispatches to the correct parser based on file extension.
        """
        file_path = Path(file_path)

        if not file_path.exists():
            raise FileNotFoundError(f"Template not found: {file_path}")

        ext = file_path.suffix.lower()
        if ext not in SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file type: {ext}. Supported: {SUPPORTED_EXTENSIONS}")

        if ext == ".docx":
            return self._parse_docx(file_path)
        elif ext == ".txt":
            return self._parse_txt(file_path)
        elif ext == ".pdf":
            return self._parse_pdf(file_path)
        else:
            raise ValueError(f"No parser implemented for: {ext}")

    def parse_directory(self, directory: Path) -> list[ParsedDocument]:
        """
        Parse all supported documents in a directory.
        Skips files that fail to parse (logs warning instead of crashing).
        """
        directory = Path(directory)
        docs = []
        for ext in SUPPORTED_EXTENSIONS:
            for file_path in sorted(directory.glob(f"*{ext}")):
                try:
                    doc = self.parse(file_path)
                    docs.append(doc)
                except Exception as e:
                    print(f"  ⚠ Skipping {file_path.name}: {e}")
        return docs

    # ── Private parsers ──────────────────────────────────────────────────

    def _parse_docx(self, file_path: Path) -> ParsedDocument:
        """Extract text + structured metadata from a .docx file."""
        docx = DocxDocument(str(file_path))

        # Collect all paragraph texts
        paragraphs = [p.text.strip() for p in docx.paragraphs if p.text.strip()]
        full_text = "\n".join(paragraphs)

        # Extract title (first non-empty paragraph)
        title = paragraphs[0] if paragraphs else file_path.stem

        # Extract fields using regex patterns
        fields = self._extract_fields(full_text)

        # Extract economic terms block (lines between "Economic Terms" and "Confirmation Terms")
        economic_terms = self._extract_section_fields(
            full_text, start_marker="Economic Terms", end_marker="Confirmation Terms"
        )

        # Extract confirmation body (narrative paragraph)
        confirmation_body = self._extract_confirmation_body(full_text)

        return ParsedDocument(
            file_path=str(file_path),
            filename=file_path.name,
            full_text=full_text,
            title=title,
            trade_type=fields.get("trade_type", self._infer_trade_type(full_text)),
            counterparty=fields.get("counterparty", "unknown"),
            jurisdiction=fields.get("jurisdiction", "unknown"),
            product=fields.get("product", "unknown"),
            version=fields.get("version", "1.0"),
            status=fields.get("status", "active"),
            economic_terms=economic_terms,
            confirmation_body=confirmation_body,
        )

    def _parse_txt(self, file_path: Path) -> ParsedDocument:
        """Parse plain text fallback."""
        full_text = file_path.read_text(encoding="utf-8")
        fields = self._extract_fields(full_text)
        title = full_text.splitlines()[0].strip() if full_text else file_path.stem

        return ParsedDocument(
            file_path=str(file_path),
            filename=file_path.name,
            full_text=full_text,
            title=title,
            trade_type=fields.get("trade_type", "unknown"),
            counterparty=fields.get("counterparty", "unknown"),
            jurisdiction=fields.get("jurisdiction", "unknown"),
            product=fields.get("product", "unknown"),
            version=fields.get("version", "1.0"),
            status=fields.get("status", "active"),
        )

    def _parse_pdf(self, file_path: Path) -> ParsedDocument:
        """
        PDF parser — requires pdfplumber.
        Install: pip install pdfplumber
        """
        try:
            import pdfplumber
        except ImportError:
            raise ImportError(
                "pdfplumber is required to parse PDFs. "
                "Run: pip install pdfplumber"
            )

        with pdfplumber.open(str(file_path)) as pdf:
            pages_text = [page.extract_text() or "" for page in pdf.pages]

        full_text = "\n".join(pages_text)
        fields = self._extract_fields(full_text)
        title = full_text.splitlines()[0].strip() if full_text else file_path.stem

        return ParsedDocument(
            file_path=str(file_path),
            filename=file_path.name,
            full_text=full_text,
            title=title,
            trade_type=fields.get("trade_type", "unknown"),
            counterparty=fields.get("counterparty", "unknown"),
            jurisdiction=fields.get("jurisdiction", "unknown"),
            product=fields.get("product", "unknown"),
            version=fields.get("version", "1.0"),
            status=fields.get("status", "active"),
        )

    # ── Helpers ──────────────────────────────────────────────────────────

    def _extract_fields(self, text: str) -> dict:
        """Run all regex patterns against the full text and return matched values."""
        results = {}
        for field_name, pattern in self._FIELD_PATTERNS.items():
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                results[field_name] = match.group(1).strip()
        return results

    def _extract_section_fields(
        self, text: str, start_marker: str, end_marker: str
    ) -> dict:
        """Extract key:value pairs from a named section of the document."""
        # Find the section block
        pattern = rf"{re.escape(start_marker)}(.*?){re.escape(end_marker)}"
        match = re.search(pattern, text, re.DOTALL | re.IGNORECASE)
        if not match:
            return {}

        section_text = match.group(1)
        fields = {}
        for line in section_text.splitlines():
            if ":" in line:
                key, _, value = line.partition(":")
                key = key.strip()
                value = value.strip()
                if key and value:
                    # Normalise key to snake_case
                    norm_key = re.sub(r"\s+", "_", key.lower())
                    fields[norm_key] = value
        return fields

    def _extract_confirmation_body(self, text: str) -> str:
        """
        Extract the main narrative paragraph of the confirmation.
        Heuristic: longest paragraph that is not a field label line.
        """
        paragraphs = [
            p.strip() for p in text.split("\n")
            if len(p.strip()) > 150  # substantial paragraphs only
        ]
        if not paragraphs:
            return ""
        # Return the longest paragraph as the main body
        return max(paragraphs, key=len)

    def _infer_trade_type(self, text: str) -> str:
        """Fallback trade type inference from document text keywords."""
        text_lower = text.lower()
        if "interest rate swap" in text_lower:
            return "Interest Rate Swap"
        if "credit default swap" in text_lower or "cds" in text_lower:
            return "Credit Default Swap"
        if "foreign exchange" in text_lower or "fx forward" in text_lower:
            return "FX Forward"
        if "equity swap" in text_lower:
            return "Equity Swap"
        if "total return swap" in text_lower:
            return "Total Return Swap"
        return "unknown"


if __name__ == "__main__":
    from src.config import TEMPLATES_DIR

    parser = DocumentParser()
    docs = parser.parse_directory(TEMPLATES_DIR)

    print(f"\n📋 Parsed {len(docs)} documents:\n")
    for doc in docs:
        print(f"  {doc.filename}")
        print(f"    Trade Type:   {doc.trade_type}")
        print(f"    Counterparty: {doc.counterparty}")
        print(f"    Jurisdiction: {doc.jurisdiction}")
        print(f"    Product:      {doc.product}")
        print(f"    Status:       {doc.status}")
        print()
